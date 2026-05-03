from __future__ import annotations

import re
from pathlib import Path

from ..config import TESSERACT_CMD


FIELD_PATTERNS: dict[str, list[str]] = {
    "rent": [r"租金[^\n]{0,20}?(\d+[元块](?:/月)?)"],
    "deposit": [r"押金[^\n]{0,20}?(\d+[元块])"],
    "lease_term": [r"租期[^\n]{0,40}?(\d{4}[./-]\d{1,2}[./-]\d{1,2}[^\n]{0,20}\d{4}[./-]\d{1,2}[./-]\d{1,2})"],
    "payment_cycle": [r"(月付|季付|半年付|年付)"],
    "breach_liability": [r"(违约[^。\n]{0,40})"],
    "termination_clause": [r"(提前退租[^。\n]{0,40}|解除合同[^。\n]{0,40})"],
    "repair_responsibility": [r"(维修[^。\n]{0,40}|修缮[^。\n]{0,40})"],
}


def extract_text(file_path: Path) -> tuple[str, list[str]]:
    suffix = file_path.suffix.lower()
    notes: list[str] = []

    if suffix in {".txt", ".md"}:
        return _read_text_file(file_path), notes

    if suffix == ".docx":
        return _read_docx(file_path), notes

    if suffix == ".pdf":
        text = _read_pdf(file_path)
        if not text.strip():
            ocr_text = _ocr_pdf(file_path, notes)
            text = ocr_text
        if not text.strip():
            notes.append("未能从 PDF 中提取到有效文本，请检查是否为扫描件或图片质量过低。")
        return text, notes

    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp"}:
        text = _ocr_image(file_path, notes)
        if not text.strip():
            notes.append("未能从图片中识别出有效文本，请检查图片清晰度或 OCR 环境。")
        return text, notes

    return _read_text_file(file_path), notes


def split_into_chunks(text: str) -> list[dict[str, str]]:
    normalized = re.sub(r"\r\n?", "\n", text).strip()
    if not normalized:
        return []

    raw_parts = re.split(r"(?=第[一二三四五六七八九十百0-9]+条)", normalized)
    if len(raw_parts) <= 1:
        raw_parts = re.split(r"\n{2,}", normalized)

    chunks: list[dict[str, str]] = []
    for idx, part in enumerate(raw_parts, start=1):
        cleaned = part.strip()
        if not cleaned:
            continue
        title = cleaned.splitlines()[0][:30]
        chunks.append(
            {
                "chunk_id": f"clause_{idx:02d}",
                "title": title,
                "text": cleaned,
            }
        )
    return chunks


def extract_fields(text: str) -> dict[str, str]:
    fields: dict[str, str] = {}
    compact = re.sub(r"[ \t]+", "", text)
    for field_name, patterns in FIELD_PATTERNS.items():
        for pattern in patterns:
            match = re.search(pattern, compact, flags=re.IGNORECASE)
            if match:
                fields[field_name] = match.group(1)
                break
    return fields


def parse_document(file_path: Path) -> dict[str, object]:
    text, notes = extract_text(file_path)
    chunks = split_into_chunks(text)
    fields = extract_fields(text)
    return {
        "text": text,
        "chunks": chunks,
        "fields": fields,
        "notes": notes,
    }


def _read_docx(file_path: Path) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(file_path))
        parts: list[str] = []
        parts.extend(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return _read_text_file(file_path)


def _read_pdf(file_path: Path) -> str:
    try:
        import fitz  # type: ignore

        with fitz.open(file_path) as doc:
            pages = []
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    pages.append(page_text)
            return "\n".join(pages)
    except Exception:
        pass

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""


def _read_text_file(file_path: Path) -> str:
    for encoding in ("utf-8", "gbk", "utf-16"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _ocr_pdf(file_path: Path, notes: list[str]) -> str:
    try:
        import fitz  # type: ignore
    except Exception:
        notes.append("未安装 PyMuPDF，无法对扫描版 PDF 做 OCR。")
        return ""

    texts: list[str] = []
    with fitz.open(file_path) as doc:
        for page in doc:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            text = _ocr_bytes(pixmap.tobytes("png"), notes)
            if text.strip():
                texts.append(text)
    return "\n".join(texts)


def _ocr_image(file_path: Path, notes: list[str]) -> str:
    return _ocr_bytes(file_path.read_bytes(), notes)


def _ocr_bytes(image_bytes: bytes, notes: list[str]) -> str:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        notes.append("未安装 OCR 依赖，无法对图片执行文字识别。")
        return ""

    if TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

    try:
        from io import BytesIO

        image = Image.open(BytesIO(image_bytes))
        return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
    except Exception as exc:
        notes.append(f"OCR 执行失败：{exc}")
        return ""
